"""Generate 5 lab-06 documentation files for PFMS."""
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

STUDENT = "Маліновський Владислав Михайлович"
TEACHER = "Гололобов Дмитро Олександрович"
UNIV    = "НАЦІОНАЛЬНИЙ УНІВЕРСИТЕТ «ЧЕРНІГІВСЬКА ПОЛІТЕХНІКА»"
DEPT    = "Кафедра комп'ютерних наук та інформаційних технологій"
YEAR    = "2026"

# ── helpers ────────────────────────────────────────────────────────────────────

def set_margins(doc):
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(1.0)

def add_page_number(doc):
    """Add page number in bottom-right footer, skip first page."""
    for i, sec in enumerate(doc.sections):
        sec.different_first_page_header_footer = True
        footer = sec.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rPr.append(rFonts)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rPr.append(sz)

def para(doc, text="", bold=False, italic=False, size=12,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(18)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rPr.insert(0, rFonts)
    return p

def heading(doc, text, level=1):
    sizes  = {1: 14, 2: 13, 3: 12}
    spaces = {1: 12, 2: 6, 3: 0}
    p = para(doc, text, bold=True, size=sizes.get(level, 12),
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=spaces.get(level, 0), space_after=6)
    return p

def h2(doc, text):
    return para(doc, text, bold=True, size=12,
                align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)

def h3(doc, text):
    return para(doc, text, bold=True, italic=True, size=12,
                align=WD_ALIGN_PARAGRAPH.LEFT, space_before=3, space_after=3)

def body(doc, text):
    return para(doc, text, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=0, space_after=6)

def bullet(doc, text):
    p = body(doc, f"– {text}")
    p.paragraph_format.left_indent = Cm(1)
    return p

def table_label(doc, num, title):
    p = para(doc, f"Таблиця {num}", bold=False, size=12,
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=12, space_after=0)
    p2 = para(doc, title, bold=True, size=12,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=3)
    return p2

def fig_caption(doc, num, title):
    p = para(doc, f"Рис. {num}. {title}", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=12)
    return p

def listing(doc, num, caption, code):
    """Add a code listing block with caption."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Courier New")
    rFonts.set(qn("w:hAnsi"), "Courier New")
    rPr.insert(0, rFonts)
    cap = para(doc, f"Лістинг {num}. {caption}", size=11,
               align=WD_ALIGN_PARAGRAPH.CENTER, space_before=3, space_after=9)
    return cap

def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Cm(w)
    return t

def title_page(doc, lab_num, lab_title, doc_type):
    set_margins(doc)
    add_page_number(doc)
    para(doc, UNIV, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, DEPT, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=60)
    para(doc, "ЗВІТ", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, f"до лабораторної роботи №{lab_num}", bold=True, size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "з навчальної дисципліни", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "«Конструювання та документування програмного забезпечення»",
         bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    para(doc, f"Тема: «{lab_title}»", bold=True, size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, f"Документ: {doc_type}", italic=True, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    # right-side info block
    p = para(doc, "", space_after=0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info = [
        ("Виконав:", f"студент гр. 123-B\n{STUDENT}"),
        ("Перевірив:", TEACHER),
    ]
    for label, val in info:
        pl = para(doc, f"{label} {val}", size=12,
                  align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
    para(doc, "", space_after=40)
    para(doc, f"Чернігів – {YEAR}", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    doc.add_page_break()

# ── DOCUMENT 1: Requirements ───────────────────────────────────────────────────

def make_requirements():
    doc = Document()
    title_page(doc, "6", "Документування програмного забезпечення",
               "Документація вимог")

    heading(doc, "1. Документація вимог PFMS")
    h2(doc, "1.1 Загальний опис системи")
    body(doc, "PFMS (Personnel File Management System) — це консольний застосунок для управління особовими справами працівників підприємства. Система розроблена на платформі .NET 10 з використанням мови програмування C# та призначена для кадрових служб і адміністраторів.")
    body(doc, "Система надає такі можливості: реєстрацію і ведення особових справ, управління статусами працівників, зберігання документів, відпусток і атестацій, повнотекстовий пошук, аудит змін та збереження даних у форматі JSON.")

    h2(doc, "1.2 Функціональні вимоги")
    body(doc, "Функціональні вимоги системи наведено в табл. 1.1.")
    table_label(doc, "1.1", "Функціональні вимоги системи")
    fr_rows = [
        ("FR-001", "Реєстрація працівника", "Система дозволяє реєструвати нового працівника із зазначенням ПІБ, відділу, посади та дати прийому.", "Обов'язкова"),
        ("FR-002", "Автоматичне створення особової справи", "При реєстрації автоматично створюється особова справа з унікальним номером.", "Обов'язкова"),
        ("FR-003", "Перегляд особової справи", "Перегляд всіх даних: документи, відпустки, атестації.", "Обов'язкова"),
        ("FR-004", "Редагування даних працівника", "Зміна контактних даних, відділу та посади.", "Обов'язкова"),
        ("FR-005", "Управління документами", "Додавання посвідчення особи, освітніх та трудових документів.", "Обов'язкова"),
        ("FR-006", "Журнал аудиту", "Автоматична реєстрація всіх змін з відміткою часу та ID користувача.", "Обов'язкова"),
        ("FR-007", "Пошук працівників", "Пошук за ПІБ, назвою відділу або посадою (часткове збіг).", "Обов'язкова"),
        ("FR-008", "Управління статусом", "Зміна статусу: Active → OnLeave, Suspended, Terminated, Archived.", "Обов'язкова"),
        ("FR-009", "Збереження та відновлення даних", "Серіалізація бази даних у JSON та десеріалізація при запуску.", "Обов'язкова"),
        ("FR-010", "Атестація та відпустки", "Додавання записів про відпустки та оцінки за результатами атестації.", "Бажана"),
    ]
    add_table(doc,
              ["ID", "Назва", "Опис", "Пріоритет"],
              fr_rows,
              widths=[2.0, 3.5, 9.0, 2.5])

    h2(doc, "1.3 Нефункціональні вимоги")
    body(doc, "Нефункціональні вимоги (табл. 1.2) визначають якісні характеристики системи.")
    table_label(doc, "1.2", "Нефункціональні вимоги")
    nfr_rows = [
        ("NFR-001", "Продуктивність",  "Час відгуку на будь-яку команду меню < 200 мс при кількості записів до 1 000 працівників."),
        ("NFR-002", "Надійність",      "При пошкодженому JSON-файлі система завантажується з seed-даних без аварійного завершення."),
        ("NFR-003", "Безпека",         "Операції «delete» та «write» доступні лише ролям Administrator та HrManager."),
        ("NFR-004", "Супроводжуваність", "Код розбито на шари (Models / Services / Data / UI) відповідно до принципу розділення відповідальностей."),
        ("NFR-005", "Переносимість",   "Застосунок компілюється та виконується на Windows, Linux і macOS (.NET 10 runtime)."),
        ("NFR-006", "Тестованість",    "Бізнес-логіка виокремлена в сервіси; покриття unit-тестами ≥ 14 тест-кейсів."),
    ]
    add_table(doc,
              ["ID", "Категорія", "Опис"],
              nfr_rows,
              widths=[2.0, 3.5, 11.5])

    h2(doc, "1.4 Варіанти використання")
    body(doc, "Основні актори системи: HR-менеджер (основний користувач) та Адміністратор (розширені права). Сценарії показано у табл. 1.3.")
    table_label(doc, "1.3", "Основні варіанти використання")
    uc_rows = [
        ("UC-01", "HR-менеджер", "Зареєструвати працівника",   "FR-001, FR-002"),
        ("UC-02", "HR-менеджер", "Переглянути особову справу", "FR-003, FR-005"),
        ("UC-03", "HR-менеджер", "Відправити у відпустку",     "FR-008"),
        ("UC-04", "HR-менеджер", "Провести пошук",             "FR-007"),
        ("UC-05", "HR-менеджер", "Додати запис відпустки",     "FR-010"),
        ("UC-06", "Адміністратор", "Звільнити та архівувати",  "FR-008"),
        ("UC-07", "Адміністратор", "Переглянути аудит",        "FR-006"),
        ("UC-08", "Система",     "Автозбереження при виході",  "FR-009"),
    ]
    add_table(doc,
              ["ID", "Актор", "Сценарій", "Вимоги"],
              uc_rows,
              widths=[2.0, 3.5, 7.5, 4.0])

    h2(doc, "1.5 Вимоги до даних")
    body(doc, "Основні сутності системи та їх атрибути наведено у табл. 1.4.")
    table_label(doc, "1.4", "Основні сутності та атрибути")
    data_rows = [
        ("Employee",         "Id, FirstName, LastName, MiddleName, BirthDate, HireDate, Status, DepartmentId, PositionId"),
        ("PersonalFile",     "FileNumber, EmployeeId, CreatedAt, UpdatedAt, IsArchived, Documents, LeaveRecords, PerformanceReviews"),
        ("Document",         "Id, Title, IssuedDate, FilePath (+підтипи: IdentityDocument, EducationDocument, ContractDocument)"),
        ("LeaveRecord",      "Id, StartDate, EndDate, Type (Annual/Sick/Unpaid), Approved"),
        ("PerformanceReview","Id, Period, Score (0–5), ReviewedBy, Comments"),
        ("Department",       "Id, Name, Code, HeadEmployeeId"),
        ("Position",         "Id, Title, SalaryGrade (1–15)"),
        ("User",             "Id, Username, Role (Administrator/HrManager), Email"),
        ("AuditLog",         "Id, Action, Timestamp, EntityId, UserId"),
    ]
    add_table(doc,
              ["Сутність", "Атрибути"],
              data_rows,
              widths=[4.5, 12.5])

    doc.save("lab6_requirements.docx")
    print("✔ lab6_requirements.docx")

# ── DOCUMENT 2: Architecture ───────────────────────────────────────────────────

def make_architecture():
    doc = Document()
    title_page(doc, "6", "Документування програмного забезпечення",
               "Документація архітектури")

    heading(doc, "2. Документація архітектури PFMS")
    h2(doc, "2.1 Загальний огляд")
    body(doc, "PFMS побудовано за шаровою (layered) архітектурою: кожен шар відповідає лише за свою область відповідальності та взаємодіє лише із сусіднім шаром. Таке розділення забезпечує тестованість, супроводжуваність та незалежну заміну реалізацій.")
    body(doc, "Склад шарів наведено у табл. 2.1.")
    table_label(doc, "2.1", "Шари архітектури PFMS")
    layers = [
        ("UI",       "PFMS.UI",       "ConsoleMenu",                          "Відображення меню, введення/виведення даних у консоль."),
        ("Services", "PFMS.Services", "EmployeeService, AuditService, LeaveService, ReviewService", "Бізнес-логіка, валідація, управління станом."),
        ("Models",   "PFMS.Models",   "Employee, PersonalFile, Document і ін.", "Доменні об'єкти, бізнес-правила (методи GetGrade, GetDuration тощо)."),
        ("Data",     "PFMS.Data",     "AppDatabase",                           "Зберігання колекцій об'єктів в пам'яті, серіалізація/десеріалізація JSON."),
        ("Enums",    "PFMS.Enums",    "EmployeeStatus, LeaveType, UserRole",   "Перелічення для типобезпечного представлення станів."),
    ]
    add_table(doc,
              ["Шар", "Простір імен", "Ключові типи", "Відповідальність"],
              layers,
              widths=[2.0, 3.5, 5.5, 6.0])

    h2(doc, "2.2 Структурна діаграма компонентів")
    body(doc, "На рис. 2.1 показано взаємодію компонентів системи.")
    # placeholder box
    p = para(doc, "[Рисунок: UML-діаграма компонентів PFMS]", italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    fig_caption(doc, "2.1", "Структурна діаграма компонентів PFMS")

    h2(doc, "2.3 Опис компонентів")
    h3(doc, "2.3.1 AppDatabase")
    body(doc, "AppDatabase — центральне сховище даних в оперативній пам'яті. Містить колекції List<T> для всіх доменних об'єктів. Методи Save() / Load() виконують серіалізацію через System.Text.Json. При пошкодженому або відсутньому файлі метод Load() повертає seed-дані без виключення.")

    h3(doc, "2.3.2 EmployeeService")
    body(doc, "EmployeeService реалізує всі операції над працівниками: Register, Update, Search, SendOnLeave, ReturnFromLeave, Suspend, Reinstate, Terminate. Кожна зміна статусу записується в журнал аудиту через AuditService. Метод Register() також автоматично генерує особову справу з номером PF-{рік}-{порядковий номер}.")

    h3(doc, "2.3.3 ConsoleMenu")
    body(doc, "ConsoleMenu реалізує інтерфейс користувача на основі числового меню. Кожен пункт меню відображає підменю або запитує введення, передає параметри до відповідного сервісу та відображає результат. ConsoleMenu залежить виключно від EmployeeService та не звертається до AppDatabase напряму.")

    h2(doc, "2.4 Діаграма класів (доменні моделі)")
    body(doc, "Ієрархія класів для Document реалізована з використанням абстрактного базового класу (рис. 2.2).")
    para(doc, "[Рисунок: UML-діаграма класів Document та підтипів]", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    fig_caption(doc, "2.2", "Ієрархія класів Document")
    body(doc, "Абстрактний клас Document визначає інтерфейс Validate() та GetDocumentType(). Конкретні підкласи: IdentityDocument (перевірка ExpiryDate), EducationDocument (перевірка IssuedDate ≤ сьогодні), ContractDocument (перевірка ContractEndDate).")
    body(doc, "Для підтримки поліморфної JSON-серіалізації використовуються атрибути [JsonPolymorphic] та [JsonDerivedType], що додають дискримінатор «$type» до кожного запису.")

    h2(doc, "2.5 Машина станів працівника")
    body(doc, "Стан працівника (EmployeeStatus) підпорядковується машині станів (рис. 2.3).")
    para(doc, "[Рисунок: State Machine діаграма EmployeeStatus]", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    fig_caption(doc, "2.3", "Машина станів EmployeeStatus")
    table_label(doc, "2.2", "Дозволені переходи статусу")
    state_rows = [
        ("New",        "Active",      "Register()"),
        ("Active",     "OnLeave",     "SendOnLeave()"),
        ("Active",     "Suspended",   "Suspend()"),
        ("Active",     "Terminated",  "Terminate() → потім Archived"),
        ("OnLeave",    "Active",      "ReturnFromLeave()"),
        ("Suspended",  "Active",      "Reinstate()"),
        ("Terminated", "Archived",    "Terminate() (автоматично)"),
    ]
    add_table(doc,
              ["Поточний стан", "Наступний стан", "Метод"],
              state_rows,
              widths=[4.5, 4.5, 8.0])

    h2(doc, "2.4 Лістинг: клас AppDatabase")
    body(doc, "Клас AppDatabase є центральним сховищем у пам'яті та відповідає за серіалізацію/десеріалізацію даних (лістинг 2.1).")
    listing(doc, "2.1", "Клас AppDatabase (Data/AppDatabase.cs)",
"""public class AppDatabase
{
    public List<Employee>          Employees        { get; set; } = new();
    public List<PersonalFile>      PersonalFiles    { get; set; } = new();
    public List<Department>        Departments      { get; set; } = new();
    public List<Position>          Positions        { get; set; } = new();
    public List<User>              Users            { get; set; } = new();
    public List<AuditLog>          AuditLogs        { get; set; } = new();

    public int NextEmployeeId  { get; set; } = 1;
    public int NextFileSeq     { get; set; } = 1;

    private static readonly JsonSerializerOptions _jsonOptions = new()
    {
        WriteIndented = true,
        PropertyNameCaseInsensitive = true
    };

    public void Save(string path)
    {
        Directory.CreateDirectory(Path.GetDirectoryName(path) ?? ".");
        File.WriteAllText(path, JsonSerializer.Serialize(this, _jsonOptions));
    }

    public static AppDatabase Load(string path)
    {
        if (!File.Exists(path))
            return CreateWithSeedData();
        try
        {
            var json = File.ReadAllText(path);
            return JsonSerializer.Deserialize<AppDatabase>(json, _jsonOptions)
                   ?? CreateWithSeedData();
        }
        catch (JsonException ex)
        {
            Console.WriteLine($"[ПОПЕРЕДЖЕННЯ] {ex.Message}");
            return CreateWithSeedData();
        }
    }
}""")

    h2(doc, "2.5 Лістинг: ієрархія класів Document")
    body(doc, "Абстрактний клас Document з поліморфною JSON-серіалізацією показано у лістингу 2.2.")
    listing(doc, "2.2", "Ієрархія класу Document (Models/Document.cs)",
"""[JsonPolymorphic(TypeDiscriminatorPropertyName = "$type")]
[JsonDerivedType(typeof(IdentityDocument),  "identity")]
[JsonDerivedType(typeof(EducationDocument), "education")]
[JsonDerivedType(typeof(ContractDocument),  "contract")]
public abstract class Document
{
    public int      Id         { get; set; }
    public string   Title      { get; set; } = string.Empty;
    public DateOnly IssuedDate { get; set; }
    public string   FilePath   { get; set; } = string.Empty;

    public abstract bool   Validate();
    public abstract string GetDocumentType();
}

public class IdentityDocument : Document
{
    public string   SeriesNumber { get; set; } = string.Empty;
    public DateOnly ExpiryDate   { get; set; }

    public override bool   Validate()         => ExpiryDate >= DateOnly.FromDateTime(DateTime.Today);
    public override string GetDocumentType()  => "Документ, що посвідчує особу";
}

public class EducationDocument : Document
{
    public string Institution { get; set; } = string.Empty;
    public string Specialty   { get; set; } = string.Empty;

    public override bool   Validate()         => IssuedDate <= DateOnly.FromDateTime(DateTime.Today);
    public override string GetDocumentType()  => "Освітній документ";
}

public class ContractDocument : Document
{
    public DateOnly? ContractEndDate { get; set; }

    public override bool   Validate()         =>
        ContractEndDate is null || ContractEndDate >= DateOnly.FromDateTime(DateTime.Today);
    public override string GetDocumentType()  => "Трудовий договір";
}""")

    h2(doc, "2.6 Формат збереження даних")
    body(doc, "Дані зберігаються в одному JSON-файлі (data/pfms.json). Файл містить всі колекції AppDatabase у форматі об'єкта JSON. Приклад структури файлу наведено у лістингу 2.3.")
    listing(doc, "2.3", "Скорочений приклад файлу pfms.json",
"""{
  "Employees": [
    { "Id": 1, "FirstName": "Тарас", "LastName": "Шевченко",
      "Status": "Active", "DepartmentId": 1, "PositionId": 1 }
  ],
  "PersonalFiles": [
    { "FileNumber": "PF-2024-1042", "EmployeeId": 1,
      "IsArchived": false, "Documents": [ ... ] }
  ],
  "NextEmployeeId": 3,
  "NextFileSeq": 1043
}""")

    doc.save("lab6_architecture.docx")
    print("✔ lab6_architecture.docx")

# ── DOCUMENT 3: Technical / Specification ──────────────────────────────────────

def make_technical():
    doc = Document()
    title_page(doc, "6", "Документування програмного забезпечення",
               "Технічна документація (специфікація)")

    heading(doc, "3. Технічна документація PFMS")
    h2(doc, "3.1 Технічний стек")
    body(doc, "Технічні характеристики розробленої системи наведено у табл. 3.1.")
    table_label(doc, "3.1", "Технічний стек PFMS")
    stack_rows = [
        ("Мова програмування",   "C# 13"),
        ("Платформа виконання",  ".NET 10.0"),
        ("Серіалізація",         "System.Text.Json (вбудований)"),
        ("Тестовий фреймворк",   "xUnit 2.9.3"),
        ("Збірка / CI",          "dotnet CLI (dotnet build, dotnet run, dotnet test)"),
        ("Формат даних",         "JSON (UTF-8)"),
        ("IDE",                  "Visual Studio Code + C# Dev Kit"),
        ("Система контролю версій", "Git"),
    ]
    add_table(doc, ["Компонент", "Значення"], stack_rows, widths=[6.0, 11.0])

    h2(doc, "3.2 Лістинг: клас Employee")
    body(doc, "Доменний клас Employee з бізнес-методами показано у лістингу 3.1.")
    listing(doc, "3.1", "Клас Employee (Models/Employee.cs)",
"""public class Employee
{
    public int            Id           { get; set; }
    public string         FirstName    { get; set; } = string.Empty;
    public string         LastName     { get; set; } = string.Empty;
    public string         MiddleName   { get; set; } = string.Empty;
    public DateOnly       BirthDate    { get; set; }
    public DateOnly       HireDate     { get; set; }
    public EmployeeStatus Status       { get; set; } = EmployeeStatus.New;
    public int            DepartmentId { get; set; }
    public int            PositionId   { get; set; }

    public string GetFullName() =>
        $"{LastName} {FirstName} {MiddleName}".Trim();

    public bool IsActive() => Status == EmployeeStatus.Active;

    public void Terminate()
    {
        if (Status is EmployeeStatus.Terminated or EmployeeStatus.Archived)
            throw new InvalidOperationException(
                $"Неможливо звільнити працівника зі статусом '{Status}'.");
        Status = EmployeeStatus.Terminated;
    }
}""")

    h2(doc, "3.3 Специфікація публічного API сервісів")
    h3(doc, "3.3.1 EmployeeService")
    body(doc, "Публічний інтерфейс EmployeeService описано у табл. 3.2.")
    table_label(doc, "3.2", "Методи EmployeeService")
    svc_rows = [
        ("Register(Employee emp, int userId)", "Employee",            "Реєструє працівника, присвоює ID та статус Active, створює PersonalFile. Кидає ArgumentException якщо FirstName або LastName порожні."),
        ("GetAll()",                           "IEnumerable<Employee>","Повертає всіх не-архівованих працівників."),
        ("GetById(int id)",                    "Employee?",           "Повертає працівника за ID або null."),
        ("Search(string query)",               "IEnumerable<Employee>","Пошук за ПІБ, назвою відділу або посади (Contains, без урахування регістру)."),
        ("Update(Employee updated, int userId)","void",               "Оновлює поля існуючого запису. Кидає KeyNotFoundException якщо ID не знайдено."),
        ("SendOnLeave(int id, int userId)",    "void",                "Active → OnLeave. Кидає InvalidOperationException якщо не Active."),
        ("ReturnFromLeave(int id, int userId)","void",                "OnLeave → Active. Кидає InvalidOperationException якщо не OnLeave."),
        ("Suspend(int id, int userId)",        "void",                "Active → Suspended."),
        ("Reinstate(int id, int userId)",      "void",                "Suspended → Active."),
        ("Terminate(int id, int userId)",      "void",                "Будь-який → Terminated → Archived; архівує PersonalFile."),
    ]
    add_table(doc,
              ["Сигнатура методу", "Повертає", "Опис / виключення"],
              svc_rows,
              widths=[5.5, 2.5, 9.0])

    h3(doc, "3.3.2 AuditService")
    body(doc, "AuditService надає один метод: Log(string action, int entityId, int userId) → void (лістинг 3.3). Метод додає запис AuditLog до колекції AppDatabase.AuditLogs із поточним UTC-часом.")
    listing(doc, "3.3", "AuditService.Log() (Services/Services.cs)",
"""public void Log(string action, int entityId, int userId)
{
    _db.AuditLogs.Add(new AuditLog
    {
        Id        = _db.NextAuditId++,
        Action    = action,
        Timestamp = DateTime.UtcNow,
        EntityId  = entityId,
        UserId    = userId
    });
}""")

    h2(doc, "3.4 Лістинги ключових методів")
    h3(doc, "3.4.1 EmployeeService.Register()")
    body(doc, "Метод Register() перевіряє вхідні дані, присвоює ID, змінює статус та створює особову справу (лістинг 3.4).")
    listing(doc, "3.4", "EmployeeService.Register() (Services/EmployeeService.cs)",
"""public Employee Register(Employee employee, int currentUserId)
{
    if (string.IsNullOrWhiteSpace(employee.FirstName))
        throw new ArgumentException("Ім'я працівника не може бути порожнім.");
    if (string.IsNullOrWhiteSpace(employee.LastName))
        throw new ArgumentException("Прізвище працівника не може бути порожнім.");

    employee.Id     = _db.NextEmployeeId++;
    employee.Status = EmployeeStatus.Active;
    _db.Employees.Add(employee);

    var file = new PersonalFile
    {
        FileNumber = GenerateFileNumber(),
        EmployeeId = employee.Id,
        CreatedAt  = DateTime.Now,
        UpdatedAt  = DateTime.Now,
    };
    _db.PersonalFiles.Add(file);

    _audit.Log("CREATE_EMPLOYEE", employee.Id, currentUserId);
    return employee;
}

private string GenerateFileNumber() =>
    $"PF-{DateTime.Now.Year}-{_db.NextFileSeq++:D4}";""")

    h3(doc, "3.4.2 PerformanceReview.GetGrade()")
    body(doc, "Метод GetGrade() реалізовано через switch-вираз із перевіркою у порядку спадання (лістинг 3.5). Порядок важливий: кожен рядок перевіряється зверху вниз, тому ≥ 4.5 перевіряється першим.")
    listing(doc, "3.5", "PerformanceReview.GetGrade() (Models/Models.cs)",
"""public string GetGrade() => Score switch
{
    >= 4.5f => "Відмінно",
    >= 3.5f => "Добре",
    >= 2.5f => "Задовільно",
    _       => "Незадовільно"
};

public bool IsPassing() => Score >= 2.5f;""")

    h3(doc, "3.4.3 PersonalFile.Archive()")
    body(doc, "Метод Archive() встановлює прапорець IsArchived та перешкоджає повторному архівуванню (лістинг 3.6).")
    listing(doc, "3.6", "PersonalFile.Archive() (Models/PersonalFile.cs)",
"""public void Archive()
{
    if (IsArchived)
        throw new InvalidOperationException(
            $"Справа {FileNumber} вже є архівованою.");

    IsArchived = true;
    UpdatedAt  = DateTime.Now;
}""")

    h3(doc, "3.4.4 EmployeeService.Search()")
    body(doc, "Метод Search() виконує пошук за частковим збігом у ПІБ, назві відділу та назві посади (лістинг 3.7).")
    listing(doc, "3.7", "EmployeeService.Search() (Services/EmployeeService.cs)",
"""public IEnumerable<Employee> Search(string query)
{
    var q = query.ToLower();
    return _db.Employees.Where(e =>
        e.GetFullName().ToLower().Contains(q)            ||
        GetDepartmentName(e.DepartmentId).ToLower().Contains(q) ||
        GetPositionTitle(e.PositionId).ToLower().Contains(q));
}""")

    h2(doc, "3.5 Специфікація моделей (зведена)")
    h3(doc, "3.5.1 Employee")
    body(doc, "Бізнес-методи класу Employee:")
    bullet(doc, "GetFullName() → string — повертає «Прізвище Ім'я По-батькові» (обрізає пробіли).")
    bullet(doc, "IsActive() → bool — true якщо Status == Active.")
    bullet(doc, "Terminate() → void — переводить у Terminated; кидає InvalidOperationException якщо вже Terminated або Archived.")

    h3(doc, "3.5.2 LeaveRecord")
    body(doc, "GetDuration() → int обчислює кількість днів включно: EndDate.DayNumber − StartDate.DayNumber + 1.")

    h3(doc, "3.5.3 PersonalFile")
    body(doc, "Archive() → void встановлює IsArchived = true. Повторний виклик кидає InvalidOperationException.")
    body(doc, "Клас зберігає колекції: Documents (List<Document>), LeaveRecords (List<LeaveRecord>), PerformanceReviews (List<PerformanceReview>).")

    h2(doc, "3.6 Обробка помилок та виключень")
    body(doc, "Стратегія обробки виключень наведена у табл. 3.3.")
    table_label(doc, "3.3", "Стратегія обробки виключень")
    exc_rows = [
        ("ArgumentException",        "Порожнє FirstName або LastName при Register().",  "Перехоплюється в ConsoleMenu, виводиться повідомлення."),
        ("KeyNotFoundException",      "ID не знайдено в GetById/Update.",               "Перехоплюється в ConsoleMenu, виводиться повідомлення."),
        ("InvalidOperationException", "Недопустимий перехід статусу; архівування вже архівованого.", "Перехоплюється в ConsoleMenu, виводиться повідомлення."),
        ("JsonException",             "Пошкоджений файл даних при Load().",             "Перехоплюється всередині AppDatabase.Load(); повертаються seed-дані."),
    ]
    add_table(doc,
              ["Тип виключення", "Причина", "Обробка"],
              exc_rows,
              widths=[4.5, 6.0, 6.5])

    h2(doc, "3.7 Конфігурація та розгортання")
    body(doc, "Шлях до файлу даних задається константою у Program.cs (рядок: const string DbPath = \"data/pfms.json\"). Змінювати шлях можна безпосередньо в коді перед компіляцією.")
    body(doc, "Для запуску необхідно:")
    bullet(doc, "Встановлений .NET 10 SDK або Runtime.")
    bullet(doc, "Виконати команду: dotnet run --project PFMS у директорії lab_04.")
    bullet(doc, "Файл даних створюється автоматично при першому збереженні.")
    body(doc, "Для запуску тестів: dotnet test у директорії lab_05/PFMS.Tests.")

    doc.save("lab6_technical.docx")
    print("✔ lab6_technical.docx")

# ── DOCUMENT 4: User Manual ────────────────────────────────────────────────────

def make_user_manual():
    doc = Document()
    title_page(doc, "6", "Документування програмного забезпечення",
               "Посібник користувача")

    heading(doc, "4. Посібник користувача PFMS")
    h2(doc, "4.1 Вступ")
    body(doc, "Цей посібник описує використання системи PFMS (Personnel File Management System) — застосунку для управління особовими справами працівників. Посібник розрахований на кадрових менеджерів та адміністраторів підприємства, які мають базові навички роботи з командним рядком.")

    h2(doc, "4.2 Системні вимоги")
    body(doc, "Мінімальні вимоги для запуску системи наведено у табл. 4.1.")
    table_label(doc, "4.1", "Системні вимоги")
    req_rows = [
        ("Операційна система", "Windows 10/11, Ubuntu 20.04+, macOS 12+"),
        ("Середовище виконання", ".NET 10 Runtime (або SDK)"),
        ("Оперативна пам'ять", "Мінімум 128 МБ"),
        ("Дисковий простір", "Мінімум 50 МБ (код + дані)"),
        ("Термінал", "Будь-який термінал з підтримкою UTF-8"),
    ]
    add_table(doc, ["Компонент", "Вимога"], req_rows, widths=[5.0, 12.0])

    h2(doc, "4.3 Запуск системи")
    body(doc, "Для запуску PFMS відкрийте термінал, перейдіть до директорії з проєктом та виконайте команду:")
    para(doc, "dotnet run --project PFMS", size=11,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_before=3, space_after=3)
    body(doc, "Після запуску система відображає головне меню (рис. 4.1).")
    para(doc, "[Рисунок: скріншот головного меню PFMS]", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    fig_caption(doc, "4.1", "Головне меню системи PFMS")

    h2(doc, "4.4 Реєстрація нового працівника")
    body(doc, "Для реєстрації нового працівника:")
    bullet(doc, "Оберіть пункт «1. Зареєструвати працівника» у головному меню.")
    bullet(doc, "Введіть прізвище, ім'я та по батькові (поле по батькові можна залишити порожнім).")
    bullet(doc, "Введіть дату народження у форматі РРРР-ММ-ДД.")
    bullet(doc, "Оберіть відділ та посаду зі списку.")
    body(doc, "Після підтвердження система присвоює унікальний ID та автоматично відкриває особову справу з номером у форматі PF-{рік}-{порядковий номер} (рис. 4.2).")
    para(doc, "[Рисунок: скріншот реєстрації нового працівника]", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    fig_caption(doc, "4.2", "Реєстрація нового працівника")

    h2(doc, "4.5 Перегляд особової справи")
    body(doc, "Оберіть пункт «3. Переглянути особову справу» та введіть ID працівника. Система відображає повну інформацію: особові дані, прикріплені документи, записи про відпустки та результати атестацій (рис. 4.3).")
    para(doc, "[Рисунок: скріншот перегляду особової справи]", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    fig_caption(doc, "4.3", "Перегляд особової справи працівника")

    h2(doc, "4.6 Пошук працівників")
    body(doc, "Оберіть пункт «6. Пошук» та введіть рядок запиту. Система шукає збіг у ПІБ, назві відділу та назві посади (без урахування регістру, часткове збіг). Результати відображаються у вигляді нумерованого списку (рис. 4.4).")
    para(doc, "[Рисунок: скріншот результатів пошуку]", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    fig_caption(doc, "4.4", "Результати пошуку")

    h2(doc, "4.7 Управління статусом працівника")
    body(doc, "Зміна статусу виконується через пункт «4. Управління статусом». Система відображає поточний статус та перелік доступних переходів. Допустимі переходи залежать від поточного статусу (як описано у розд. 2.5 архітектурної документації). Недопустимий перехід виводить повідомлення про помилку без завершення роботи програми.")

    h2(doc, "4.8 Збереження та відновлення даних")
    body(doc, "Дані зберігаються автоматично при виборі пункту «0. Зберегти та вийти». Файл data/pfms.json створюється або перезаписується у директорії запуску. При наступному запуску система автоматично завантажує збережені дані. Якщо файл пошкоджений або відсутній, система стартує з демонстраційними даними без повідомлення про помилку (за винятком попередження у консолі).")

    h2(doc, "4.9 Журнал аудиту")
    body(doc, "Усі операції створення, оновлення та зміни статусу автоматично фіксуються в журналі аудиту. Для перегляду оберіть пункт «7. Журнал аудиту» у головному меню. Кожен запис містить: дату і час (UTC), тип дії, ID сутності та ID користувача.")

    h2(doc, "4.10 Типові помилки та їх усунення")
    body(doc, "Найпоширеніші проблеми та способи їх вирішення наведено у табл. 4.2.")
    table_label(doc, "4.2", "Типові помилки та їх усунення")
    err_rows = [
        ("«Ім'я не може бути порожнім»",       "Введено порожнє поле FirstName або LastName.",           "Введіть ім'я та прізвище заново."),
        ("«Працівника з ID=X не знайдено»",    "Введено неіснуючий ID.",                                "Перевірте ID через пункт «2. Список працівників»."),
        ("«Операція недоступна для статусу X»","Спроба зміни статусу у неправильній послідовності.",    "Перевірте поточний статус і виберіть допустимий перехід."),
        ("«[ПОПЕРЕДЖЕННЯ] Не вдалося завантажити дані»", "Файл pfms.json пошкоджений.",               "Видаліть або відновіть файл; система завантажиться з demo-даних."),
    ]
    add_table(doc,
              ["Повідомлення", "Причина", "Дія"],
              err_rows,
              widths=[5.0, 5.5, 6.5])

    doc.save("lab6_user_manual.docx")
    print("✔ lab6_user_manual.docx")

# ── DOCUMENT 5: Marketing ──────────────────────────────────────────────────────

def make_marketing():
    doc = Document()
    title_page(doc, "6", "Документування програмного забезпечення",
               "Маркетингова документація")

    heading(doc, "5. Маркетингова документація PFMS")
    h2(doc, "5.1 Опис продукту")
    body(doc, "PFMS (Personnel File Management System) — програмний застосунок для ведення особових справ персоналу. Система автоматизує рутинні кадрові процеси: облік працівників, управління документами, відстеження відпусток і атестацій, а також забезпечує повний журнал аудиту всіх операцій.")
    body(doc, "PFMS є автономним, кросплатформним рішенням, яке не потребує підключення до бази даних або мережі Інтернет: усі дані зберігаються локально у JSON-файлі.")

    h2(doc, "5.2 Цільова аудиторія")
    body(doc, "Продукт орієнтований на такі категорії споживачів (табл. 5.1).")
    table_label(doc, "5.1", "Цільова аудиторія PFMS")
    aud_rows = [
        ("Малий бізнес (до 50 ос.)",  "Підприємства, яким потрібна просте та безкоштовне рішення для обліку кадрів без хмарних підписок."),
        ("HR-відділи середніх підприємств", "Організації, що використовують внутрішні інструменти автоматизації без зовнішніх залежностей."),
        ("Навчальні заклади",          "Кафедри та лабораторії, де PFMS виступає навчальним прикладом розробки кадрових ІС."),
        ("Розробники-початківці",      "Студенти та junior-розробники, що вивчають архітектуру багатошарових C#-застосунків."),
    ]
    add_table(doc,
              ["Сегмент", "Потреба / варіант використання"],
              aud_rows,
              widths=[5.0, 12.0])

    h2(doc, "5.3 Ключові функції та переваги")
    body(doc, "Переваги PFMS порівняно з ручним обліком та базовими таблицями наведено у табл. 5.2.")
    table_label(doc, "5.2", "Ключові переваги PFMS")
    feat_rows = [
        ("Автоматична нумерація справ",    "Унікальний номер PF-{рік}-{N} присвоюється без участі оператора."),
        ("Поліморфна структура документів","Один файл даних зберігає паспорти, дипломи та трудові договори з різними атрибутами."),
        ("Машина станів",                  "Чітко визначені допустимі переходи між статусами: Active, OnLeave, Suspended, Archived."),
        ("Журнал аудиту",                  "Кожна зміна фіксується з відміткою часу та ID відповідального."),
        ("Відновлення після збою",         "Пошкоджений файл даних не призводить до краху — система стартує з демо-даними."),
        ("Кросплатформність",              "Працює на Windows, Linux, macOS без змін у коді."),
        ("Відкритий вихідний код",         "Повна прозорість; можливість кастомізації під конкретне підприємство."),
    ]
    add_table(doc,
              ["Функція / перевага", "Деталі"],
              feat_rows,
              widths=[5.5, 11.5])

    h2(doc, "5.4 Конкурентне порівняння")
    body(doc, "Порівняння PFMS з альтернативними підходами наведено у табл. 5.3.")
    table_label(doc, "5.3", "Конкурентне порівняння")
    comp_rows = [
        ("Excel-таблиці",          "Безкоштовно; знайомий інтерфейс",    "Немає аудиту; ручна нумерація; схильний до помилок"),
        ("Хмарні HR-системи",      "Багатий функціонал; веб-доступ",     "Дорога підписка; залежність від Інтернету; витік даних"),
        ("1C: Підприємство",        "Локальна система; широкий функціонал", "Висока вартість; складне налаштування; залежність від постачальника"),
        ("PFMS (наш продукт)",     "Безкоштовно; кросплатформно; аудит; автовідновлення", "Консольний інтерфейс; одночасна робота одного користувача"),
    ]
    add_table(doc,
              ["Рішення", "Переваги", "Недоліки"],
              comp_rows,
              widths=[4.0, 7.0, 6.0])

    h2(doc, "5.5 Умови розповсюдження та підтримки")
    body(doc, "PFMS поширюється як відкритий навчальний проєкт. Умови використання:")
    bullet(doc, "Безкоштовне використання у навчальних та некомерційних цілях.")
    bullet(doc, "Вихідний код доступний для перегляду, форкування та модифікації.")
    bullet(doc, "Технічна підтримка здійснюється через систему контролю версій (Git Issues).")
    bullet(doc, "Відповідальність за результати комерційного використання покладається на користувача.")

    h2(doc, "5.6 Дорожня карта розвитку")
    body(doc, "Планований розвиток продукту (табл. 5.4).")
    table_label(doc, "5.4", "Дорожня карта PFMS")
    road_rows = [
        ("v1.0 (поточна)", "Консольний інтерфейс, JSON-сховище, повна бізнес-логіка, 14 unit-тестів."),
        ("v1.1",           "Веб-інтерфейс (Blazor або React) замість консольного меню."),
        ("v1.2",           "Підтримка SQLite як альтернативного сховища."),
        ("v2.0",           "Багатокористувацький режим з автентифікацією та REST API."),
    ]
    add_table(doc,
              ["Версія", "Зміни / функції"],
              road_rows,
              widths=[3.0, 14.0])

    doc.save("lab6_marketing.docx")
    print("✔ lab6_marketing.docx")

# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import os
    os.chdir(r"d:\Education\software-engineering-labs\software-construction-documentation\lab_06")
    make_requirements()
    make_architecture()
    make_technical()
    make_user_manual()
    make_marketing()
    print("Всі документи згенеровано у папці lab_06/")
