"""Generate single lab6_pfms.docx — combined report with all 5 documentation types."""
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

STUDENT = "Маліновський Владислав Михайлович"
TEACHER = "Гололобов Дмитро Олександрович"
UNIV    = "НАЦІОНАЛЬНИЙ УНІВЕРСИТЕТ «ЧЕРНІГІВСЬКА ПОЛІТЕХНІКА»"
DEPT    = "Кафедра комп'ютерних наук та інформаційних технологій"
YEAR    = "2026"

# ── helpers ────────────────────────────────────────────────────────────────────

def set_margins(doc):
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(1.0)

def add_page_number(doc):
    for sec in doc.sections:
        sec.different_first_page_header_footer = True
        footer = sec.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        for tag, text in [("w:fldChar", None), ("w:instrText", "PAGE"), ("w:fldChar", None)]:
            el = OxmlElement(tag)
            if tag == "w:fldChar":
                el.set(qn("w:fldCharType"), "begin" if text is None and not run._r.findall(qn("w:fldChar")) else "end")
            elif text:
                el.text = text
            run._r.append(el)
        fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText"); instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
        run._r.clear()
        run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rPr.append(rFonts)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rPr.append(sz)

def tnr_run(run, size=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.insert(0, rFonts)

def p(doc, text="", bold=False, italic=False, size=12,
      align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6, indent=None):
    para = doc.add_paragraph()
    para.paragraph_format.alignment    = align
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    para.paragraph_format.line_spacing = Pt(18)
    if indent is not None:
        para.paragraph_format.left_indent = Cm(indent)
    if text:
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        tnr_run(run, size)
    return para

def h1(doc, text):
    return p(doc, text, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, sb=12, sa=6)

def h2(doc, text):
    return p(doc, text, bold=True, size=13, align=WD_ALIGN_PARAGRAPH.LEFT, sb=10, sa=4)

def h3(doc, text):
    return p(doc, text, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=3)

def body(doc, text):
    return p(doc, text, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6)

def bullet(doc, text):
    return p(doc, f"– {text}", size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=3, indent=1.0)

def tbl_label(doc, num, title):
    p(doc, f"Таблиця {num}", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, sb=10, sa=0)
    p(doc, title, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=3)

def listing(doc, num, caption, code):
    para = doc.add_paragraph()
    para.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(0)
    para.paragraph_format.left_indent  = Cm(0.5)
    para.paragraph_format.line_spacing = Pt(14)
    run = para.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Courier New")
    rFonts.set(qn("w:hAnsi"), "Courier New")
    rPr.insert(0, rFonts)
    p(doc, f"Лістинг {num}. {caption}", size=11,
      align=WD_ALIGN_PARAGRAPH.CENTER, sb=3, sa=9)

def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        tnr_run(run, 11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            tnr_run(run, 11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                if ci < len(widths):
                    row.cells[ci].width = Cm(w)
    return t

# ── Title page ─────────────────────────────────────────────────────────────────

def title_page(doc):
    set_margins(doc)
    add_page_number(doc)
    p(doc, UNIV, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
    p(doc, DEPT, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=60)
    p(doc, "ЗВІТ", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
    p(doc, "до лабораторної роботи №6", bold=True, size=14,
      align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
    p(doc, "з навчальної дисципліни", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
    p(doc, "«Конструювання та документування програмного забезпечення»",
      bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
    p(doc, "Тема: «Документування програмного забезпечення»",
      bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, sa=40)
    p(doc, f"Виконав: студент гр. 123-B\n{STUDENT}", size=12,
      align=WD_ALIGN_PARAGRAPH.RIGHT, sa=6)
    p(doc, f"Перевірив: {TEACHER}", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, sa=40)
    p(doc, f"Чернігів – {YEAR}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
    doc.add_page_break()

# ── Section 1: Task ────────────────────────────────────────────────────────────

def section_task(doc):
    h1(doc, "1. Завдання")
    body(doc, "Тема роботи: Документування програмного забезпечення.")
    body(doc, "Мета роботи: Розробити документацію програмного забезпечення відповідно до індивідуального завдання.")
    h2(doc, "Практична частина")
    body(doc, "Розробити документацію розроблюваного застосунку (PFMS — система управління особовими справами), що включає:")
    bullet(doc, "документацію вимог;")
    bullet(doc, "документацію з архітектури;")
    bullet(doc, "технічну документацію (специфікацію);")
    bullet(doc, "документацію користувача (посібник користувача);")
    bullet(doc, "маркетингову документацію.")

# ── Section 2: Requirements ────────────────────────────────────────────────────

def section_requirements(doc):
    h1(doc, "2. Документація вимог")
    h2(doc, "2.1 Загальний опис системи")
    body(doc, "PFMS (Personnel File Management System) — консольний застосунок для управління особовими справами працівників підприємства. Розроблено на платформі .NET 10, мова C#. Система призначена для кадрових менеджерів та адміністраторів.")
    body(doc, "Основні можливості: реєстрація та ведення особових справ, управління статусами (машина станів), зберігання документів різних типів, запис відпусток та атестацій, повнотекстовий пошук, журнал аудиту, JSON-серіалізація.")

    h2(doc, "2.2 Функціональні вимоги")
    body(doc, "Функціональні вимоги до системи наведено у табл. 2.1.")
    tbl_label(doc, "2.1", "Функціональні вимоги системи")
    add_table(doc,
        ["ID", "Назва", "Опис", "Пріоритет"],
        [
            ("FR-001", "Реєстрація працівника",           "Реєстрація нового працівника з ПІБ, відділом, посадою.", "Обов'язкова"),
            ("FR-002", "Створення особової справи",       "Автоматичне відкриття справи PF-{рік}-{N} при реєстрації.", "Обов'язкова"),
            ("FR-003", "Перегляд особової справи",        "Повний перегляд: документи, відпустки, атестації.", "Обов'язкова"),
            ("FR-004", "Редагування даних",               "Зміна ПІБ, відділу, посади існуючого працівника.", "Обов'язкова"),
            ("FR-005", "Управління документами",          "Додавання посвідчень особи, освітніх і трудових договорів.", "Обов'язкова"),
            ("FR-006", "Журнал аудиту",                   "Автоматичний запис усіх змін з UTC-часом та ID користувача.", "Обов'язкова"),
            ("FR-007", "Пошук",                           "Пошук за ПІБ, назвою відділу або посади (часткове збіг).", "Обов'язкова"),
            ("FR-008", "Управління статусом",             "Переходи: Active↔OnLeave, Active↔Suspended, Active→Terminated.", "Обов'язкова"),
            ("FR-009", "Збереження та відновлення даних", "Серіалізація у JSON; відновлення з seed-даних при збої.", "Обов'язкова"),
            ("FR-010", "Відпустки та атестації",          "Додавання LeaveRecord та PerformanceReview до особової справи.", "Бажана"),
        ],
        widths=[1.8, 3.5, 9.2, 2.5])

    h2(doc, "2.3 Нефункціональні вимоги")
    body(doc, "Нефункціональні вимоги наведено у табл. 2.2.")
    tbl_label(doc, "2.2", "Нефункціональні вимоги")
    add_table(doc,
        ["ID", "Категорія", "Вимога"],
        [
            ("NFR-001", "Продуктивність",    "Час відгуку < 200 мс при ≤ 1 000 записів."),
            ("NFR-002", "Надійність",        "Пошкоджений JSON не призводить до краш-зупинки."),
            ("NFR-003", "Безпека",           "Операції write/delete доступні лише ролям Administrator та HrManager."),
            ("NFR-004", "Супроводжуваність", "Чіткий поділ на шари: Models / Services / Data / UI."),
            ("NFR-005", "Переносимість",     "Запуск на Windows, Linux, macOS (.NET 10 Runtime)."),
            ("NFR-006", "Тестованість",      "Бізнес-логіка ізольована в сервісах; ≥ 14 unit-тестів."),
        ],
        widths=[2.0, 3.5, 11.5])

    h2(doc, "2.4 Варіанти використання")
    body(doc, "Основні варіанти використання системи наведено у табл. 2.3.")
    tbl_label(doc, "2.3", "Варіанти використання")
    add_table(doc,
        ["ID", "Актор", "Сценарій", "Вимоги"],
        [
            ("UC-01", "HR-менеджер",   "Зареєструвати нового працівника",    "FR-001, FR-002"),
            ("UC-02", "HR-менеджер",   "Переглянути та оновити особову справу", "FR-003, FR-004"),
            ("UC-03", "HR-менеджер",   "Відправити у відпустку / повернути",  "FR-008"),
            ("UC-04", "HR-менеджер",   "Виконати пошук за ПІБ або відділом", "FR-007"),
            ("UC-05", "HR-менеджер",   "Додати відпустку / атестацію",       "FR-010"),
            ("UC-06", "Адміністратор", "Звільнити та архівувати справу",      "FR-008"),
            ("UC-07", "Адміністратор", "Переглянути журнал аудиту",           "FR-006"),
            ("UC-08", "Система",       "Автозбереження даних у JSON при виході", "FR-009"),
        ],
        widths=[2.0, 3.5, 7.5, 4.0])

    h2(doc, "2.5 Сутності та атрибути")
    body(doc, "Перелік основних сутностей системи наведено у табл. 2.4.")
    tbl_label(doc, "2.4", "Основні сутності та атрибути")
    add_table(doc,
        ["Сутність", "Ключові атрибути"],
        [
            ("Employee",          "Id, FirstName, LastName, MiddleName, BirthDate, HireDate, Status, DepartmentId, PositionId"),
            ("PersonalFile",      "FileNumber, EmployeeId, CreatedAt, UpdatedAt, IsArchived, Documents, LeaveRecords, PerformanceReviews"),
            ("Document",          "Id, Title, IssuedDate, FilePath + підтипи: IdentityDocument, EducationDocument, ContractDocument"),
            ("LeaveRecord",       "Id, StartDate, EndDate, Type (Annual/Sick/Unpaid), Approved"),
            ("PerformanceReview", "Id, Period, Score (1.0–5.0), ReviewedBy, Comments"),
            ("Department",        "Id, Name, Code, HeadEmployeeId"),
            ("Position",          "Id, Title, SalaryGrade (1–15)"),
            ("User",              "Id, Username, Role (Administrator/HrManager), Email"),
            ("AuditLog",          "Id, Action, Timestamp (UTC), EntityId, UserId"),
        ],
        widths=[4.5, 12.5])

# ── Section 3: Architecture ────────────────────────────────────────────────────

def section_architecture(doc):
    h1(doc, "3. Документація архітектури")
    h2(doc, "3.1 Шарова архітектура")
    body(doc, "PFMS реалізує шарову архітектуру: кожен шар відповідає лише за свою область та взаємодіє лише із сусіднім шаром (табл. 3.1).")
    tbl_label(doc, "3.1", "Шари архітектури PFMS")
    add_table(doc,
        ["Шар", "Простір імен", "Ключові типи", "Відповідальність"],
        [
            ("UI",       "PFMS.UI",       "ConsoleMenu",                                   "Консольне меню, введення/виведення."),
            ("Services", "PFMS.Services", "EmployeeService, AuditService, PersonalFileService", "Бізнес-логіка, валідація, управління станом."),
            ("Models",   "PFMS.Models",   "Employee, PersonalFile, Document і ін.",        "Доменні об'єкти, бізнес-правила."),
            ("Data",     "PFMS.Data",     "AppDatabase",                                   "Зберігання колекцій, серіалізація JSON."),
            ("Enums",    "PFMS.Enums",    "EmployeeStatus, LeaveType, UserRole",            "Типобезпечні перелічення станів."),
        ],
        widths=[2.0, 3.5, 5.5, 6.0])

    h2(doc, "3.2 Машина станів EmployeeStatus")
    body(doc, "Стан працівника підпорядковується машині станів. Дозволені переходи наведено у табл. 3.2.")
    tbl_label(doc, "3.2", "Дозволені переходи статусу")
    add_table(doc,
        ["Поточний стан", "Наступний стан", "Метод EmployeeService"],
        [
            ("New",        "Active",     "Register()"),
            ("Active",     "OnLeave",    "SendOnLeave()"),
            ("Active",     "Suspended",  "Suspend()"),
            ("Active",     "Terminated", "Terminate() → потім Archived"),
            ("OnLeave",    "Active",     "ReturnFromLeave()"),
            ("Suspended",  "Active",     "Reinstate()"),
            ("Terminated", "Archived",   "Terminate() (автоматично)"),
        ],
        widths=[4.5, 4.5, 8.0])

    h2(doc, "3.3 Лістинг: клас AppDatabase")
    body(doc, "Центральне сховище з методами Save() і Load() показано у лістингу 3.1.")
    listing(doc, "3.1", "AppDatabase (Data/AppDatabase.cs)",
"""public class AppDatabase
{
    public List<Employee>     Employees     { get; set; } = new();
    public List<PersonalFile> PersonalFiles { get; set; } = new();
    public List<Department>   Departments   { get; set; } = new();
    public List<Position>     Positions     { get; set; } = new();
    public List<User>         Users         { get; set; } = new();
    public List<AuditLog>     AuditLogs     { get; set; } = new();

    public int NextEmployeeId { get; set; } = 1;
    public int NextFileSeq    { get; set; } = 1;

    private static readonly JsonSerializerOptions _jsonOptions = new()
    {
        WriteIndented = true,
        PropertyNameCaseInsensitive = true
    };

    public void Save(string path)
    {
        Directory.CreateDirectory(Path.GetDirectoryName(path) ?? ".");
        File.WriteAllText(path, JsonSerializer.Serialize(this, _jsonOptions));
    }

    public static AppDatabase Load(string path)
    {
        if (!File.Exists(path))
            return CreateWithSeedData();
        try
        {
            var json = File.ReadAllText(path);
            return JsonSerializer.Deserialize<AppDatabase>(json, _jsonOptions)
                   ?? CreateWithSeedData();
        }
        catch (JsonException ex)
        {
            Console.WriteLine($"[ПОПЕРЕДЖЕННЯ] {ex.Message}");
            return CreateWithSeedData();
        }
    }
}""")

    h2(doc, "3.4 Лістинг: ієрархія класів Document")
    body(doc, "Поліморфна JSON-серіалізація реалізована через атрибути [JsonPolymorphic] та [JsonDerivedType] (лістинг 3.2).")
    listing(doc, "3.2", "Ієрархія Document (Models/Document.cs)",
"""[JsonPolymorphic(TypeDiscriminatorPropertyName = "$type")]
[JsonDerivedType(typeof(IdentityDocument),  "identity")]
[JsonDerivedType(typeof(EducationDocument), "education")]
[JsonDerivedType(typeof(ContractDocument),  "contract")]
public abstract class Document
{
    public int      Id         { get; set; }
    public string   Title      { get; set; } = string.Empty;
    public DateOnly IssuedDate { get; set; }
    public string   FilePath   { get; set; } = string.Empty;

    public abstract bool   Validate();
    public abstract string GetDocumentType();
}

public class IdentityDocument : Document
{
    public string   SeriesNumber { get; set; } = string.Empty;
    public DateOnly ExpiryDate   { get; set; }
    public override bool   Validate()        => ExpiryDate >= DateOnly.FromDateTime(DateTime.Today);
    public override string GetDocumentType() => "Документ, що посвідчує особу";
}

public class EducationDocument : Document
{
    public string Institution { get; set; } = string.Empty;
    public string Specialty   { get; set; } = string.Empty;
    public override bool   Validate()        => IssuedDate <= DateOnly.FromDateTime(DateTime.Today);
    public override string GetDocumentType() => "Освітній документ";
}

public class ContractDocument : Document
{
    public DateOnly? ContractEndDate { get; set; }
    public override bool   Validate()        =>
        ContractEndDate is null || ContractEndDate >= DateOnly.FromDateTime(DateTime.Today);
    public override string GetDocumentType() => "Трудовий договір";
}""")

# ── Section 4: Technical Specification ────────────────────────────────────────

def section_technical(doc):
    h1(doc, "4. Технічна документація (специфікація)")
    h2(doc, "4.1 Технічний стек")
    body(doc, "Технічні характеристики розробленої системи наведено у табл. 4.1.")
    tbl_label(doc, "4.1", "Технічний стек PFMS")
    add_table(doc,
        ["Компонент", "Значення"],
        [
            ("Мова програмування",      "C# 13"),
            ("Платформа",               ".NET 10.0"),
            ("Серіалізація",            "System.Text.Json"),
            ("Тестовий фреймворк",      "xUnit 2.9.3"),
            ("Збірка",                  "dotnet CLI"),
            ("Формат даних",            "JSON (UTF-8)"),
            ("Система версій",          "Git"),
        ],
        widths=[6.0, 11.0])

    h2(doc, "4.2 Лістинг: клас Employee")
    body(doc, "Доменний клас Employee з повним переліком властивостей та бізнес-методів наведено у лістингу 4.1.")
    listing(doc, "4.1", "Employee (Models/Employee.cs)",
"""public class Employee
{
    public int            Id           { get; set; }
    public string         FirstName    { get; set; } = string.Empty;
    public string         LastName     { get; set; } = string.Empty;
    public string         MiddleName   { get; set; } = string.Empty;
    public DateOnly       BirthDate    { get; set; }
    public DateOnly       HireDate     { get; set; }
    public EmployeeStatus Status       { get; set; } = EmployeeStatus.New;
    public int            DepartmentId { get; set; }
    public int            PositionId   { get; set; }

    public string GetFullName() =>
        $"{LastName} {FirstName} {MiddleName}".Trim();

    public bool IsActive() => Status == EmployeeStatus.Active;

    public void Terminate()
    {
        if (Status is EmployeeStatus.Terminated or EmployeeStatus.Archived)
            throw new InvalidOperationException(
                $"Неможливо звільнити працівника зі статусом '{Status}'.");
        Status = EmployeeStatus.Terminated;
    }

    public override string ToString() =>
        $"[{Id:D4}] {GetFullName()} | Статус: {Status} | Прийнято: {HireDate}";
}""")

    h2(doc, "4.3 Лістинг: EmployeeService")
    body(doc, "Ключові методи EmployeeService — Register(), Search() та GetActiveEmployee() — показано у лістингу 4.2.")
    listing(doc, "4.2", "EmployeeService (Services/EmployeeService.cs)",
"""public class EmployeeService
{
    private readonly AppDatabase  _db;
    private readonly AuditService _audit;

    public EmployeeService(AppDatabase db, AuditService audit)
    {
        _db    = db;
        _audit = audit;
    }

    public Employee Register(Employee employee, int currentUserId)
    {
        if (string.IsNullOrWhiteSpace(employee.FirstName))
            throw new ArgumentException("Ім'я працівника не може бути порожнім.");
        if (string.IsNullOrWhiteSpace(employee.LastName))
            throw new ArgumentException("Прізвище працівника не може бути порожнім.");

        employee.Id     = _db.NextEmployeeId++;
        employee.Status = EmployeeStatus.Active;
        _db.Employees.Add(employee);

        var file = new PersonalFile
        {
            FileNumber = GenerateFileNumber(),
            EmployeeId = employee.Id,
            CreatedAt  = DateTime.Now,
            UpdatedAt  = DateTime.Now,
        };
        _db.PersonalFiles.Add(file);

        _audit.Log("CREATE_EMPLOYEE", employee.Id, currentUserId);
        return employee;
    }

    public IEnumerable<Employee> Search(string query)
    {
        var q = query.ToLower();
        return _db.Employees.Where(e =>
            e.GetFullName().ToLower().Contains(q)                    ||
            GetDepartmentName(e.DepartmentId).ToLower().Contains(q)  ||
            GetPositionTitle(e.PositionId).ToLower().Contains(q));
    }

    public void SendOnLeave(int employeeId, int currentUserId)
    {
        var emp = GetActiveEmployee(employeeId);
        emp.Status = EmployeeStatus.OnLeave;
        _audit.Log("STATUS_ON_LEAVE", employeeId, currentUserId);
    }

    public void Terminate(int employeeId, int currentUserId)
    {
        var emp  = GetById(employeeId)
            ?? throw new KeyNotFoundException($"Працівника з ID={employeeId} не знайдено.");
        emp.Terminate();
        var file = _db.PersonalFiles.FirstOrDefault(f => f.EmployeeId == employeeId);
        file?.Archive();
        emp.Status = EmployeeStatus.Archived;
        _audit.Log("TERMINATE_EMPLOYEE", employeeId, currentUserId);
    }

    private Employee GetActiveEmployee(int id)
    {
        var emp = GetById(id)
            ?? throw new KeyNotFoundException($"Працівника з ID={id} не знайдено.");
        if (!emp.IsActive())
            throw new InvalidOperationException(
                $"Операція недоступна для статусу '{emp.Status}'.");
        return emp;
    }

    private string GenerateFileNumber() =>
        $"PF-{DateTime.Now.Year}-{_db.NextFileSeq++:D4}";
}""")

    h2(doc, "4.4 Лістинг: AuditService")
    body(doc, "AuditService відповідає виключно за запис подій до журналу (лістинг 4.3).")
    listing(doc, "4.3", "AuditService (Services/Services.cs)",
"""public class AuditService
{
    private readonly AppDatabase _db;

    public AuditService(AppDatabase db) { _db = db; }

    public void Log(string action, int entityId, int userId)
    {
        _db.AuditLogs.Add(new AuditLog
        {
            Id        = _db.NextAuditId++,
            Action    = action,
            Timestamp = DateTime.UtcNow,
            EntityId  = entityId,
            UserId    = userId
        });
    }

    public IEnumerable<AuditLog> GetAll() =>
        _db.AuditLogs.OrderByDescending(l => l.Timestamp);
}""")

    h2(doc, "4.5 Лістинг: PerformanceReview та LeaveRecord")
    body(doc, "Бізнес-методи моделей PerformanceReview та LeaveRecord показано у лістингу 4.4.")
    listing(doc, "4.4", "PerformanceReview та LeaveRecord (Models/Models.cs)",
"""public class PerformanceReview
{
    public int    Id         { get; set; }
    public string Period     { get; set; } = string.Empty;
    public float  Score      { get; set; }
    public string ReviewedBy { get; set; } = string.Empty;
    public string Comments   { get; set; } = string.Empty;

    public string GetGrade() => Score switch
    {
        >= 4.5f => "Відмінно",
        >= 3.5f => "Добре",
        >= 2.5f => "Задовільно",
        _       => "Незадовільно"
    };

    public bool IsPassing() => Score >= 2.5f;
}

public class LeaveRecord
{
    public int      Id        { get; set; }
    public DateOnly StartDate { get; set; }
    public DateOnly EndDate   { get; set; }
    public LeaveType Type     { get; set; }
    public bool     Approved  { get; set; } = false;

    public int  GetDuration() => EndDate.DayNumber - StartDate.DayNumber + 1;
    public bool IsApproved()  => Approved;
}""")

    h2(doc, "4.6 Лістинг: PersonalFile")
    body(doc, "Клас PersonalFile зберігає колекції пов'язаних сутностей та захищає від змін після архівування (лістинг 4.5).")
    listing(doc, "4.5", "PersonalFile (Models/PersonalFile.cs)",
"""public class PersonalFile
{
    public string   FileNumber  { get; set; } = string.Empty;
    public int      EmployeeId  { get; set; }
    public DateTime CreatedAt   { get; set; } = DateTime.Now;
    public DateTime UpdatedAt   { get; set; } = DateTime.Now;
    public bool     IsArchived  { get; set; } = false;

    public List<Document>          Documents          { get; set; } = new();
    public List<LeaveRecord>       LeaveRecords       { get; set; } = new();
    public List<PerformanceReview> PerformanceReviews { get; set; } = new();

    public void Archive()
    {
        if (IsArchived)
            throw new InvalidOperationException(
                $"Справа {FileNumber} вже є архівованою.");
        IsArchived = true;
        UpdatedAt  = DateTime.Now;
    }

    public void AddDocument(Document document)
    {
        if (IsArchived)
            throw new InvalidOperationException(
                $"Справа {FileNumber} архівована і не може бути змінена.");
        Documents.Add(document);
        UpdatedAt = DateTime.Now;
    }
}""")

    h2(doc, "4.7 Специфікація методів EmployeeService")
    body(doc, "Публічний інтерфейс EmployeeService наведено у табл. 4.2.")
    tbl_label(doc, "4.2", "Методи EmployeeService")
    add_table(doc,
        ["Метод", "Повертає", "Опис / виключення"],
        [
            ("Register(emp, userId)",        "Employee",             "Реєстрація; ArgumentException якщо ім'я порожнє."),
            ("GetAll()",                     "IEnumerable<Employee>","Всі не-архівовані."),
            ("GetById(id)",                  "Employee?",            "За ID або null."),
            ("Search(query)",                "IEnumerable<Employee>","Contains по ПІБ / відділу / посаді."),
            ("Update(updated, userId)",      "void",                 "KeyNotFoundException якщо не знайдено."),
            ("SendOnLeave(id, userId)",      "void",                 "Active→OnLeave; InvalidOperationException інакше."),
            ("ReturnFromLeave(id, userId)",  "void",                 "OnLeave→Active."),
            ("Suspend(id, userId)",          "void",                 "Active→Suspended."),
            ("Reinstate(id, userId)",        "void",                 "Suspended→Active."),
            ("Terminate(id, userId)",        "void",                 "→Terminated→Archived; архівує PersonalFile."),
        ],
        widths=[5.0, 2.8, 9.2])

    h2(doc, "4.8 Обробка помилок")
    body(doc, "Стратегію обробки виключень наведено у табл. 4.3.")
    tbl_label(doc, "4.3", "Обробка виключень")
    add_table(doc,
        ["Виключення", "Причина", "Обробка"],
        [
            ("ArgumentException",        "Порожнє ім'я при Register().",               "Перехоплюється в ConsoleMenu."),
            ("KeyNotFoundException",     "ID не знайдено.",                            "Перехоплюється в ConsoleMenu."),
            ("InvalidOperationException","Недопустимий перехід статусу / архівування.", "Перехоплюється в ConsoleMenu."),
            ("JsonException",            "Пошкоджений pfms.json.",                     "Перехоплюється в AppDatabase.Load(); seed-дані."),
        ],
        widths=[4.5, 6.0, 6.5])

# ── Section 5: User Manual ────────────────────────────────────────────────────

def section_user_manual(doc):
    h1(doc, "5. Документація користувача (посібник)")
    h2(doc, "5.1 Системні вимоги")
    tbl_label(doc, "5.1", "Системні вимоги")
    add_table(doc,
        ["Компонент", "Вимога"],
        [
            ("ОС",                   "Windows 10/11, Ubuntu 20.04+, macOS 12+"),
            ("Середовище виконання", ".NET 10 Runtime або SDK"),
            ("Оперативна пам'ять",   "≥ 128 МБ"),
            ("Дисковий простір",     "≥ 50 МБ"),
            ("Термінал",             "UTF-8 сумісний"),
        ],
        widths=[5.0, 12.0])

    h2(doc, "5.2 Запуск системи")
    body(doc, "Відкрийте термінал у директорії з проєктом та виконайте:")
    listing(doc, "5.1", "Команда запуску",
"""dotnet run --project PFMS""")
    body(doc, "Після запуску відображається головне меню з переліком доступних операцій.")

    h2(doc, "5.3 Основні операції")
    body(doc, "Основні операції та їх порядок виконання наведено у табл. 5.2.")
    tbl_label(doc, "5.2", "Основні операції системи")
    add_table(doc,
        ["Операція", "Пункт меню", "Порядок виконання"],
        [
            ("Реєстрація працівника",    "1",   "Оберіть пункт → введіть ПІБ → дату народження → оберіть відділ і посаду."),
            ("Перегляд особової справи", "3",   "Оберіть пункт → введіть ID працівника."),
            ("Зміна статусу",            "4",   "Оберіть пункт → введіть ID → оберіть нову дію."),
            ("Пошук",                    "6",   "Оберіть пункт → введіть рядок запиту (мін. 1 символ)."),
            ("Журнал аудиту",            "7",   "Оберіть пункт — відображається хронологічний список подій."),
            ("Зберегти та вийти",        "0",   "Дані записуються у data/pfms.json, програма завершується."),
        ],
        widths=[4.5, 2.5, 10.0])

    h2(doc, "5.4 Типові помилки та вирішення")
    tbl_label(doc, "5.3", "Типові помилки та їх вирішення")
    add_table(doc,
        ["Повідомлення", "Причина", "Дія"],
        [
            ("«Ім'я не може бути порожнім»",       "Порожнє FirstName або LastName.",          "Введіть ім'я та прізвище."),
            ("«Працівника з ID=X не знайдено»",    "Неіснуючий ID.",                           "Перевірте ID через список працівників."),
            ("«Операція недоступна для статусу X»","Недопустимий перехід статусу.",             "Перевірте поточний статус і виберіть правильну дію."),
            ("«[ПОПЕРЕДЖЕННЯ] Не вдалося завантажити дані»", "Пошкоджений pfms.json.",         "Видаліть або відновіть файл."),
        ],
        widths=[5.0, 5.5, 6.5])

# ── Section 6: Marketing ──────────────────────────────────────────────────────

def section_marketing(doc):
    h1(doc, "6. Маркетингова документація")
    h2(doc, "6.1 Опис продукту")
    body(doc, "PFMS — програмний застосунок для автоматизації обліку особових справ персоналу. Система усуває ручне ведення документації, забезпечує цілісність даних через машину станів та фіксує всі зміни в журналі аудиту. Розповсюджується як відкритий проєкт на .NET 10.")

    h2(doc, "6.2 Цільова аудиторія")
    tbl_label(doc, "6.1", "Цільова аудиторія")
    add_table(doc,
        ["Сегмент", "Потреба"],
        [
            ("Малий бізнес (до 50 осіб)",       "Просте та безкоштовне рішення без хмарних підписок."),
            ("HR-відділи середніх підприємств",  "Внутрішній інструмент без зовнішніх залежностей."),
            ("Навчальні заклади",                "Навчальний приклад багатошарового C#-застосунку."),
            ("Розробники-початківці",            "Вивчення архітектурних патернів на реальному проєкті."),
        ],
        widths=[5.0, 12.0])

    h2(doc, "6.3 Ключові переваги")
    tbl_label(doc, "6.2", "Ключові переваги PFMS")
    add_table(doc,
        ["Функція", "Деталі"],
        [
            ("Автонумерація справ",          "PF-{рік}-{N} — унікальний номер без участі оператора."),
            ("Поліморфні документи",         "Один JSON зберігає паспорти, дипломи й договори різних типів."),
            ("Машина станів",                "Чіткі дозволені переходи унеможливлюють некоректні операції."),
            ("Повний аудит",                 "Кожна зміна зафіксована з UTC-часом і ID відповідального."),
            ("Відновлення після збою",       "Пошкоджений файл не зупиняє систему — запуск з demo-даних."),
            ("Кросплатформність",            "Windows / Linux / macOS без змін у коді."),
            ("Відкритий код",                "Повна прозорість і можливість кастомізації."),
        ],
        widths=[5.5, 11.5])

    h2(doc, "6.4 Конкурентне порівняння")
    tbl_label(doc, "6.3", "Порівняння з альтернативами")
    add_table(doc,
        ["Рішення", "Переваги", "Недоліки"],
        [
            ("Excel-таблиці",      "Безкоштовно; знайомий інтерфейс.",       "Немає аудиту; ручна нумерація; схильний до помилок."),
            ("Хмарні HR-системи",  "Багатий функціонал; веб-доступ.",        "Дорога підписка; залежність від Інтернету."),
            ("1C: Підприємство",   "Локально; широкий функціонал.",          "Висока вартість; складне налаштування."),
            ("PFMS",               "Безкоштовно; кросплатформно; аудит; автовідновлення.", "Консоль; один користувач одночасно."),
        ],
        widths=[3.5, 7.0, 6.5])

    h2(doc, "6.5 Дорожня карта")
    tbl_label(doc, "6.4", "Дорожня карта розвитку PFMS")
    add_table(doc,
        ["Версія", "Зміни"],
        [
            ("v1.0 (поточна)", "Консоль, JSON, повна бізнес-логіка, 14 unit-тестів."),
            ("v1.1",           "Веб-інтерфейс (Blazor або React)."),
            ("v1.2",           "SQLite як альтернативне сховище."),
            ("v2.0",           "Багатокористувацький режим з REST API та автентифікацією."),
        ],
        widths=[3.0, 14.0])

# ── Section 7: Conclusion ─────────────────────────────────────────────────────

def section_conclusion(doc):
    h1(doc, "7. Висновок")
    body(doc, "У ході виконання лабораторної роботи №6 було розроблено повний комплект документації для застосунку PFMS (Personnel File Management System).")
    body(doc, "Підготовлені такі типи документів:")
    bullet(doc, "Документація вимог — 10 функціональних (FR-001–FR-010) та 6 нефункціональних вимог, 8 варіантів використання, перелік сутностей та атрибутів.")
    bullet(doc, "Документація архітектури — опис 5-шарової архітектури, машина станів EmployeeStatus, лістинги AppDatabase та ієрархії Document з поліморфною серіалізацією.")
    bullet(doc, "Технічна документація — лістинги 5 ключових класів і методів (Employee, EmployeeService, AuditService, PerformanceReview, PersonalFile), специфікація публічного API, стратегія обробки помилок.")
    bullet(doc, "Посібник користувача — системні вимоги, інструкції запуску та виконання основних операцій, таблиця типових помилок та їх вирішення.")
    bullet(doc, "Маркетингова документація — опис продукту, цільова аудиторія, ключові переваги, конкурентне порівняння та дорожня карта розвитку.")
    body(doc, "Розроблена документація відповідає вимогам до оформлення: формат A4, шрифт Times New Roman 12 пт, нумерація сторінок у правому нижньому колонтитулі, таблиці та лістинги з підписами.")

# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import os
    os.chdir(r"d:\Education\software-engineering-labs\software-construction-documentation\lab_06")

    doc = Document()
    title_page(doc)
    section_task(doc)
    doc.add_page_break()
    section_requirements(doc)
    doc.add_page_break()
    section_architecture(doc)
    doc.add_page_break()
    section_technical(doc)
    doc.add_page_break()
    section_user_manual(doc)
    doc.add_page_break()
    section_marketing(doc)
    doc.add_page_break()
    section_conclusion(doc)

    doc.save("lab6_pfms.docx")
    print("✔ lab6_pfms.docx")
